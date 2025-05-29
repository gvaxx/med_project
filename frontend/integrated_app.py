import streamlit as st
import requests
import json
import os
from datetime import datetime
import tempfile
from moviepy import VideoFileClip
from pydub import AudioSegment
import io
import base64
from docx import Document
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration
import markdown
import re
from bs4 import BeautifulSoup

# Service URLs from environment variables
MEDICAL_DOC_SERVICE_URL = os.getenv("MEDICAL_DOC_SERVICE_URL", "http://medical-doc-service:8000")
AUDIO_TRANSCRIPTION_SERVICE_URL = os.getenv("AUDIO_TRANSCRIPTION_SERVICE_URL", "http://audio-transcription-service:8004")
LLM_SERVICE_URL = os.getenv("LLM_SERVICE_URL", "http://llm-service:8003")

# Configure the page
st.set_page_config(
    page_title="Медицинская документация",
    page_icon="🏥",
    layout="wide"
)

# Добавление глобальных стилей
st.markdown("""
<style>
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    h1, h2, h3, h4 {
        color: #2c3e50;
    }
    h1 {
        border-bottom: 2px solid #3498db;
        padding-bottom: 10px;
        margin-bottom: 20px;
    }
    .stTextArea > div > div > textarea {
        border-radius: 8px;
        border-color: #e0e0e0;
    }
    .stTextArea > label {
        font-weight: bold;
        color: #2c3e50;
    }
    .stAlert {
        border-radius: 8px;
    }
    div[data-testid="stExpander"] {
        border-radius: 8px;
        border: 1px solid #e0e0e0;
    }
</style>
""", unsafe_allow_html=True)

def get_available_models():
    """Get available LLM models"""
    try:
        response = requests.get(f"{LLM_SERVICE_URL}/available_models")
        if response.status_code == 200:
            return response.json()
    except Exception:
        pass
    return {"gpt-3.5-turbo": True, "gpt-4": True}  # Default models if service unreachable

# Get available LLM models
available_models = get_available_models()

def extract_audio_from_video(video_file):
    """Extract audio from video file"""
    try:
        # Create a temporary file to save the video
        with tempfile.NamedTemporaryFile(suffix='.mp4', delete=False) as temp_video:
            temp_video.write(video_file.read())
            temp_video_path = temp_video.name

        # Extract audio using moviepy
        video = VideoFileClip(temp_video_path)
        
        # Create a temporary file for the audio
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_audio:
            video.audio.write_audiofile(temp_audio.name)
            audio_segment = AudioSegment.from_wav(temp_audio.name)
            
        # Cleanup temporary files
        os.unlink(temp_video_path)
        os.unlink(temp_audio.name)
        video.close()
        
        return audio_segment
    except Exception as e:
        st.error(f"Ошибка при извлечении аудио из видео: {str(e)}")
        return None

def load_audio_file(audio_file):
    """Load audio file into AudioSegment"""
    try:
        # Read the uploaded file into memory
        audio_bytes = audio_file.read()
        
        # Create a temporary file with the correct extension
        ext = audio_file.name.split('.')[-1].lower()
        with tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False) as temp_audio:
            temp_audio.write(audio_bytes)
            audio_segment = AudioSegment.from_file(temp_audio.name, format=ext)
            os.unlink(temp_audio.name)
        return audio_segment
    except Exception as e:
        st.error(f"Ошибка при загрузке аудио файла: {str(e)}")
        return None

def combine_audio_segments(audio_segments):
    """Combine multiple audio segments into one"""
    try:
        combined = AudioSegment.empty()
        for segment in audio_segments:
            combined += segment
        return combined
    except Exception as e:
        st.error(f"Ошибка при объединении аудио файлов: {str(e)}")
        return None

def audio_segment_to_file(audio_segment):
    """Convert AudioSegment to file-like object"""
    try:
        buffer = io.BytesIO()
        audio_segment.export(buffer, format="wav")
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Ошибка при конвертации аудио: {str(e)}")
        return None 

def check_services():
    """Check if all required services are available"""
    services_status = {}
    
    # Check audio transcription service
    try:
        response = requests.get(f"{AUDIO_TRANSCRIPTION_SERVICE_URL}/test", timeout=5)
        if response.status_code == 200:
            result = response.json()
            services_status["audio"] = {
                "status": "available",
                "models": result.get("models", {})
            }
        else:
            services_status["audio"] = {"status": "error", "message": response.text}
    except Exception as e:
        services_status["audio"] = {"status": "error", "message": str(e)}
    
    # Check medical doc service
    try:
        response = requests.get(f"{MEDICAL_DOC_SERVICE_URL}/health", timeout=5)
        if response.status_code == 200:
            services_status["medical"] = {"status": "available"}
        else:
            services_status["medical"] = {"status": "error", "message": response.text}
    except Exception as e:
        services_status["medical"] = {"status": "error", "message": str(e)}
    
    return services_status

def transcribe_audio(audio_file):
    """Transcribe audio file using the transcription service"""
    try:
        files = {"file": (audio_file.name, audio_file, audio_file.type)}
        data = {
            "model_type": "rnnt",  # используем RNNT модель по умолчанию
            "long_form": "true"    # включаем поддержку длинных аудио
        }
        
        st.info(f"Отправка аудиофайла: {audio_file.name} ({audio_file.type})")
        
        response = requests.post(
            f"{AUDIO_TRANSCRIPTION_SERVICE_URL}/transcribe",
            files=files,
            data=data,
            timeout=600  # увеличиваем timeout до 10 минут для длинных аудио
        )
        
        if response.status_code == 200:
            result = response.json()
            # Проверяем формат ответа
            if "utterances" in result:
                # Для длинной транскрипции объединяем все сегменты
                return "\n".join([u["transcription"] for u in result["utterances"]])
            elif "transcription" in result:
                # Для обычной транскрипции
                return result["transcription"]
            else:
                st.error("Неожиданный формат ответа от сервиса транскрибации")
                return None
        else:
            st.error(f"Ошибка при транскрибации: {response.text}")
            return None
    except requests.exceptions.Timeout:
        st.error("Превышено время ожидания ответа от сервера. Аудиофайл может быть слишком длинным.")
        return None
    except Exception as e:
        st.error(f"Ошибка при транскрибации: {str(e)}")
        return None

def process_transcript(text, selected_model):
    """Process transcript into medical documentation"""
    try:
        # Создаем статус-контейнер для отображения прогресса
        status_container = st.empty()
        
        response = requests.post(
            f"{MEDICAL_DOC_SERVICE_URL}/process_transcript",
            json={
                "transcript": text,
                "model_type": selected_model,
                "parameters": {
                    "temperature": 0.3,
                    "max_tokens": 8000
                }
            },
            stream=True,
            timeout=600  # 10 минут таймаут
        )
        
        if response.status_code == 200:
            for line in response.iter_lines():
                if line:
                    data = json.loads(line.decode('utf-8'))
                    status = data.get("status")
                    
                    if status == "completed":
                        status_container.empty()
                        return data.get("structured_doc")
                    elif status == "error":
                        st.error(data.get("message"))
                        return None
                    else:
                        # Обновляем статус
                        with status_container:
                            st.info(data.get("message"))
        else:
            st.error(f"Ошибка при обработке транскрипта: {response.text}")
            return None
    except Exception as e:
        st.error(f"Ошибка при обработке транскрипта: {str(e)}")
        return None

def analyze_medical_doc(medical_doc, selected_model):
    """Analyze medical documentation and generate recommendations"""
    try:
        # Создаем статус-контейнер для отображения прогресса
        status_container = st.empty()
        
        response = requests.post(
            f"{MEDICAL_DOC_SERVICE_URL}/analyze",
            json={
                "medical_doc": medical_doc,
                "model_type": selected_model,
                "top_k": 3,
                "parameters": {
                    "temperature": 0.3,
                    "max_tokens": 8000
                }
            },
            stream=True,
            timeout=600  # 10 минут таймаут
        )
        
        if response.status_code == 200:
            for line in response.iter_lines():
                if line:
                    data = json.loads(line.decode('utf-8'))
                    status = data.get("status")
                    
                    if status == "completed":
                        status_container.empty()
                        return {
                            "recommendations": data.get("recommendations"),
                            "similar_documents": data.get("similar_documents")
                        }
                    elif status == "error":
                        st.error(data.get("message"))
                        return None
                    else:
                        # Обновляем статус
                        with status_container:
                            st.info(data.get("message"))
        else:
            st.error(f"Ошибка при анализе документа: {response.text}")
            return None
    except Exception as e:
        st.error(f"Ошибка при анализе документа: {str(e)}")
        return None

def create_pdf(medical_doc, recommendations=None):
    """Создание PDF документа с использованием WeasyPrint с поддержкой Markdown"""
    try:
        # Преобразуем Markdown в HTML
        medical_doc_html = markdown.markdown(medical_doc)
        recommendations_html = markdown.markdown(recommendations) if recommendations else ""
        
        # Создаем HTML контент по частям
        html_header = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Медицинская документация</title>
            <style>
                @page {
                    size: A4;
                    margin: 2cm;
                }
                body {
                    font-family: 'DejaVu Sans', sans-serif;
                    font-size: 12pt;
                }
                h1 {
                    font-size: 18pt;
                    text-align: center;
                    margin-bottom: 20px;
                }
                h2 {
                    font-size: 16pt;
                    margin-top: 30px;
                    margin-bottom: 15px;
                }
                h3 {
                    font-size: 14pt;
                    margin-top: 20px;
                    margin-bottom: 10px;
                }
                p {
                    margin-bottom: 10px;
                }
                ul, ol {
                    margin-left: 20px;
                    margin-bottom: 10px;
                }
                li {
                    margin-bottom: 5px;
                }
                a {
                    color: #0066cc;
                    text-decoration: none;
                }
                blockquote {
                    border-left: 3px solid #ccc;
                    margin-left: 0;
                    padding-left: 15px;
                    color: #555;
                }
                code {
                    font-family: monospace;
                    background-color: #f5f5f5;
                    padding: 2px 4px;
                    border-radius: 3px;
                }
                .page-break {
                    page-break-after: always;
                }
                footer {
                    position: fixed;
                    bottom: 0;
                    width: 100%;
                    text-align: center;
                    font-size: 9pt;
                }
            </style>
        </head>
        <body>
            <h1>Медицинская документация</h1>
            <div>
        """
        
        # Добавляем документацию
        doc_content = medical_doc_html
        
        # Часть с рекомендациями
        recommendations_content = ""
        if recommendations:
            recommendations_content = """
            <div class="page-break"></div>
            <h1>Рекомендации</h1>
            <div>
            """ + recommendations_html + """
            </div>
            """
        
        # Часть с футером
        current_date = datetime.now().strftime("%d.%m.%Y %H:%M")
        footer_html = f"""
            </div>
            <footer>
                <p>Документ сгенерирован {current_date}</p>
            </footer>
        </body>
        </html>
        """
        
        # Собираем весь HTML вместе
        html_content = html_header + doc_content + recommendations_content + footer_html
        
        # Конфигурация шрифтов
        font_config = FontConfiguration()
        
        # Рендеринг HTML в PDF
        pdf_buffer = io.BytesIO()
        HTML(string=html_content).write_pdf(
            pdf_buffer,
            font_config=font_config,
            stylesheets=[
                CSS(string='@page { size: A4; margin: 2cm; }')
            ]
        )
        
        pdf_buffer.seek(0)
        return pdf_buffer.read()
        
    except Exception as e:
        st.error(f"Ошибка при создании PDF: {str(e)}")
        return None

def create_docx(medical_doc, recommendations=None):
    """Создание DOCX документа с поддержкой Markdown"""
    try:
        doc = Document()
        
        # Заголовок
        doc.add_heading('Медицинская документация', 0)
        
        # Конвертация Markdown в HTML для обработки
        html_content = markdown.markdown(medical_doc)
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Функция для рекурсивной обработки HTML элементов и применения форматирования
        def process_element(element, parent):
            if element.name is None:  # Текстовый узел
                if element.string and element.string.strip():
                    parent.add_paragraph(element.string)
                return
                
            if element.name == 'h1':
                parent.add_heading(element.text, 1)
            elif element.name == 'h2':
                parent.add_heading(element.text, 2)
            elif element.name == 'h3':
                parent.add_heading(element.text, 3)
            elif element.name == 'p':
                p = parent.add_paragraph()
                for child in element.children:
                    if child.name == 'strong' or child.name == 'b':
                        p.add_run(child.text).bold = True
                    elif child.name == 'em' or child.name == 'i':
                        p.add_run(child.text).italic = True
                    elif child.name == 'code':
                        p.add_run(child.text).font.name = 'Courier New'
                    elif child.name is None:  # Простой текст
                        p.add_run(child.string)
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    p = parent.add_paragraph(style='List Bullet')
                    p.add_run(li.text)
            elif element.name == 'ol':
                for idx, li in enumerate(element.find_all('li', recursive=False), 1):
                    p = parent.add_paragraph(style='List Number')
                    p.add_run(li.text)
            elif element.name == 'blockquote':
                p = parent.add_paragraph(style='Quote')
                p.add_run(element.text)
            else:  # Для других элементов, обрабатываем детей рекурсивно
                for child in element.children:
                    process_element(child, parent)
        
        # Обработка HTML для основного документа
        # Исправление: проверяем, существует ли тег body, и если нет, обрабатываем корневые элементы напрямую
        if soup.body:
            elements_to_process = list(soup.body.children)
        else:
            elements_to_process = list(soup.children)
            
        for element in elements_to_process:
            process_element(element, doc)
        
        # Добавляем рекомендации, если они есть
        if recommendations:
            # Разрыв страницы
            doc.add_page_break()
            doc.add_heading('Рекомендации', 0)
            
            # Аналогичная обработка для рекомендаций
            html_recs = markdown.markdown(recommendations)
            soup_recs = BeautifulSoup(html_recs, 'html.parser')
            
            # Исправление: такая же проверка для рекомендаций
            if soup_recs.body:
                rec_elements = list(soup_recs.body.children)
            else:
                rec_elements = list(soup_recs.children)
                
            for element in rec_elements:
                process_element(element, doc)
        
        # Создаем байтовый объект с DOCX
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.getvalue()
    except Exception as e:
        st.error(f"Ошибка при создании DOCX: {str(e)}")
        return None

# Initialize session state
if "transcription" not in st.session_state:
    st.session_state.transcription = None
if "medical_doc" not in st.session_state:
    st.session_state.medical_doc = None
if "recommendations" not in st.session_state:
    st.session_state.recommendations = None
if "similar_documents" not in st.session_state:
    st.session_state.similar_documents = None
if "doc_edited" not in st.session_state:
    st.session_state.doc_edited = False
if "generate_doc" not in st.session_state:
    st.session_state.generate_doc = False
if "generate_rec" not in st.session_state:
    st.session_state.generate_rec = False

# Функции для установки флагов в session_state вместо прямого вызова генерации в callback
def set_generate_doc():
    st.session_state.generate_doc = True

def set_generate_rec():
    st.session_state.generate_rec = True

# Main UI
st.title("🏥 Автоматизация медицинской документации")

# Базовая инструкция без лишнего HTML
st.markdown("""
### Процесс работы:
1. Загрузите аудиозапись приема
2. Получите автоматическую транскрипцию
3. Сгенерируйте медицинскую документацию
4. Просмотрите и отредактируйте сгенерированную документацию
5. Получите рекомендации на основе документации
""")

# File upload section
st.header("1. Загрузка файлов")
uploaded_file = st.file_uploader(
    "Загрузите аудио или видео файлы",
    type=['mp3', 'wav', 'm4a', 'mp4', 'avi', 'mov'],
)

# Process files and generate documentation
if uploaded_file:
    if st.button("🎯 Начать обработку"):
        with st.spinner("Транскрибация аудио..."):
            transcription = transcribe_audio(uploaded_file)
            if transcription:
                st.session_state.transcription = transcription
                st.rerun()

# Display transcription and model selection for document generation
if st.session_state.transcription:
    st.header("2. Транскрипция")
    st.text_area(
        "Текст транскрипции",
        st.session_state.transcription,
        height=200,
        disabled=True
    )
    
    # Add model selection and document generation button
    st.header("3. Генерация медицинской документации")
    
    # Упрощенный интерфейс без лишних HTML-вставок
    doc_gen_model = st.selectbox(
        "Выберите модель ИИ для генерации документации",
        options=list(available_models.keys()),
        key="doc_model_select"
    )
    st.button("📝 Сгенерировать", 
                key="generate_doc_button", 
                on_click=set_generate_doc)
    
    # Генерация документации при установленном флаге
    if st.session_state.generate_doc:
        doc_gen_placeholder = st.empty()
        with doc_gen_placeholder.container():
            with st.spinner("⏳ Генерация медицинской документации..."):
                medical_doc = process_transcript(st.session_state.transcription, st.session_state.doc_model_select)
                if medical_doc:
                    st.session_state.medical_doc = medical_doc
                    st.session_state.doc_edited = False
                    st.session_state.generate_doc = False
                    st.rerun()
                else:
                    st.error("Не удалось сгенерировать документацию")
                    st.session_state.generate_doc = False

# Display and edit medical documentation
if st.session_state.medical_doc:
    st.header("4. Медицинская документация")
    st.markdown("*Отредактируйте документацию при необходимости:*")
    
    # Увеличиваем поле для редактирования
    edited_doc = st.text_area(
        "Медицинская документация",
        st.session_state.medical_doc,
        height=400,
        key="medical_doc_editor"
    )
    
    if edited_doc != st.session_state.medical_doc:
        st.session_state.medical_doc = edited_doc
        st.session_state.doc_edited = True
    
    # Model selection and recommendation generation
    st.header("5. Генерация рекомендаций")
    
    # Упрощенный интерфейс без лишних HTML-вставок
    rec_gen_model = st.selectbox(
        "Выберите модель ИИ для генерации рекомендаций",
        options=list(available_models.keys()),
        key="rec_model_select"
    )
    st.button("🔍 Сгенерировать", 
            key="generate_rec_button",
            on_click=set_generate_rec)
    
    # Генерация рекомендаций при установленном флаге
    if st.session_state.generate_rec:
        rec_gen_placeholder = st.empty()
        with rec_gen_placeholder.container():
            with st.spinner("⏳ Генерация рекомендаций..."):
                result = analyze_medical_doc(st.session_state.medical_doc, st.session_state.rec_model_select)
                if result:
                    st.session_state.recommendations = result["recommendations"]
                    st.session_state.similar_documents = result["similar_documents"]
                    st.session_state.generate_rec = False
                    st.rerun()
                else:
                    st.error("Не удалось сгенерировать рекомендации")
                    st.session_state.generate_rec = False

# Display and edit recommendations
if st.session_state.recommendations:
    st.header("6. Рекомендации")
    
    # Allow editing recommendations with larger field
    edited_recommendations = st.text_area(
        "Рекомендации (отредактируйте при необходимости)",
        st.session_state.recommendations,
        height=300,
        key="recommendations_editor"
    )
    
    if edited_recommendations != st.session_state.recommendations:
        st.session_state.recommendations = edited_recommendations
    
    # Display similar cases
    if st.session_state.similar_documents:
        st.header("7. Похожие случаи")
        for idx, doc in enumerate(st.session_state.similar_documents, 1):
            with st.expander(f"Случай {idx} (Схожесть: {doc['similarity']:.2f})", expanded=False):
                st.markdown(doc['content'])
                if 'metadata' in doc:
                    st.markdown("**Метаданные:**")
                    if doc['metadata'].get('specialty'):
                        st.markdown(f"- Специальность: {doc['metadata']['specialty']}")
                    if doc['metadata'].get('diagnoses'):
                        st.markdown("- Диагнозы: " + ", ".join(doc['metadata']['diagnoses']))
    
    # Document export section
    st.header("8. Экспорт документации")
    
    # Простое описание экспорта
    st.markdown("""
    ### Выберите формат для экспорта документации:
    - **DOCX** - стандартный формат для Microsoft Word
    - **PDF** - универсальный формат для просмотра на любых устройствах
    - **Markdown** - текстовый формат с разметкой
    """)
    
    # Создаем кнопки экспорта в одном ряду
    export_col1, export_col2, export_col3 = st.columns(3)
    
    # Имена файлов для экспорта
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    docx_filename = f"medical_doc_{timestamp}.docx"
    pdf_filename = f"medical_doc_{timestamp}.pdf"
    md_filename = f"medical_doc_{timestamp}.md"
    
    # Подготовка Markdown-контента
    markdown_content = f"""# Медицинская документация
    
{st.session_state.medical_doc}

# Рекомендации

{st.session_state.recommendations}

# Похожие случаи
"""
    
    # Кнопка экспорта в DOCX
    with export_col1:
        with st.spinner("Создание DOCX документа..."):
            docx_bytes = create_docx(st.session_state.medical_doc, st.session_state.recommendations)
            if docx_bytes:
                st.download_button(
                    label="📄 Скачать DOCX",
                    data=docx_bytes,
                    file_name=docx_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_docx",
                    use_container_width=True
                )
            else:
                st.error("Ошибка при создании DOCX документа")
    
    # Кнопка экспорта в PDF
    with export_col2:
        with st.spinner("Создание PDF документа..."):
            pdf_bytes = create_pdf(st.session_state.medical_doc, st.session_state.recommendations)
            if pdf_bytes:
                st.download_button(
                    label="📊 Скачать PDF",
                    data=pdf_bytes,
                    file_name=pdf_filename,
                    mime="application/pdf",
                    key="download_pdf",
                    use_container_width=True
                )
            else:
                st.error("Ошибка при создании PDF документа")
    
    # Кнопка экспорта в Markdown
    with export_col3:
        st.download_button(
            label="📝 Скачать Markdown",
            data=markdown_content,
            file_name=md_filename,
            mime="text/markdown",
            key="download_md",
            use_container_width=True
        )
